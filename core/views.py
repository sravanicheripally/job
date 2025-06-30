from django.shortcuts import render
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth import authenticate
from .forms import SignUpForm, StudentProfileForm,AppliedJobForm
from .models import Plan,JobApplier,AppliedJob,JobPosting
import razorpay
from django.views.decorators.csrf import csrf_protect
import pandas as pd
import os, re
from datetime import datetime
from django.db import connections
from django.shortcuts import render
from django.contrib import messages
from datetime import date, timedelta
from datetime import timezone
import pdfplumber
import docx
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt


# Razorpay credentials
RAZORPAY_KEY_ID='rzp_test_KcubyV1kaxj5bY'
RAZORPAY_KEY_SECRET='hKVpAnCTgFf8s48y40K58MWr'
client = razorpay.Client(auth=(RAZORPAY_KEY_ID, RAZORPAY_KEY_SECRET))

def signup(request):
    if request.method == 'POST':
        form = SignUpForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('home')
    else:
        form = SignUpForm()
    return render(request, 'signup.html', {'form': form})

@csrf_protect
def login_view(request):
    if request.method == 'POST':
        form = AuthenticationForm(data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)

            # Check if user is a JobApplier
            if hasattr(user, 'jobapplier'):
                return redirect('job_applier_dashboard')
            else:
                return redirect('home')
    else:
        form = AuthenticationForm()
    return render(request, 'login.html', {'form': form})

def job_applier_dashboard(request):
    return render(request, 'job_applier_dashboard.html')


def home(request):
    return render(request, 'home.html')

@login_required
def plans(request):
    plans = Plan.objects.all()
    return render(request, 'plans.html', {'plans': plans})


@login_required
def subscribe(request, plan_id):
    plan = get_object_or_404(Plan, id=plan_id)

    if request.method == 'POST':
        amount = int(plan.price * 100)
        order = client.order.create({
            'amount': amount,
            'currency': 'INR',
            'payment_capture': 1
        })

        context = {
            'plan': plan,
            'order_id': order['id'],
            'razorpay_key': RAZORPAY_KEY_ID,
            'amount': amount,
            'user': request.user
        }
        return render(request, 'payment_success.html', context)

    return redirect('plans')

def extract_text_from_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])



def parse_resume(text):
    data = {
        'firstname': '',
        'lastname': '',
        'skills': '',
        'email': '',
        'description': ''
    }

    lines = text.strip().split('\n')
    lines = [line.strip() for line in lines if line.strip()]  # remove empty lines

    # --- Extract Email ---
    email_match = re.search(r'[\w\.-]+@[\w\.-]+', text)
    if email_match:
        data['email'] = email_match.group(0)

    # --- Extract Name (best guess from top lines) ---
    # Option 1: All uppercase name in first few lines
    for line in lines[:5]:
        if re.match(r'^[A-Z]{2,}\s+[A-Z]{2,}$', line):  # e.g., "SRAVANI CHERIPALLY"
            name_parts = line.strip().split()
            data['firstname'] = name_parts[0].capitalize()
            data['lastname'] = name_parts[1].capitalize()
            break

    # Option 2: Fallback to capitalized words from first line
    if not data['firstname'] or not data['lastname']:
        for line in lines[:3]:
            capitalized_words = re.findall(r'\b[A-Z][a-z]{2,}\b', line)
            if len(capitalized_words) >= 2:
                data['firstname'] = capitalized_words[0]
                data['lastname'] = capitalized_words[1]
                break

    # Option 3: fallback to name from email
    if (not data['firstname'] or not data['lastname']) and data['email']:
        name_part = data['email'].split('@')[0].replace('.', ' ').replace('_', ' ')
        name_tokens = name_part.split()
        if len(name_tokens) >= 2:
            data['firstname'] = data['firstname'] or name_tokens[0].capitalize()
            data['lastname'] = data['lastname'] or name_tokens[1].capitalize()

    # --- Extract Skills ---
    skill_keywords = ['Python', 'Django', 'React', 'SQL', 'Machine Learning', 'Java', 'C++', 'HTML', 'CSS']
    skills_found = [skill for skill in skill_keywords if skill.lower() in text.lower()]
    data['skills'] = ", ".join(skills_found)

    # --- Extract Description (first 5 non-empty lines after name) ---
    data['description'] = "\n".join(lines[1:6])

    return data


@csrf_exempt
@login_required
def parse_resume_api(request):
    if request.method == 'POST' and request.FILES.get('resume'):
        resume = request.FILES['resume']
        text = ""
        if resume.name.endswith('.pdf'):
            with pdfplumber.open(resume) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        elif resume.name.endswith('.docx'):
            doc = docx.Document(resume)
            text = "\n".join([para.text for para in doc.paragraphs])

        parsed = parse_resume(text)
        return JsonResponse(parsed)

    return JsonResponse({'error': 'Invalid request'}, status=400)



@login_required
def profile_form(request):
    initial_data = {}

    if request.method == 'POST':
        form = StudentProfileForm(request.POST, request.FILES)

        if request.FILES.get('resume'):
            resume = request.FILES['resume']
            text = ""
            if resume.name.endswith('.pdf'):
                text = extract_text_from_pdf(resume)
            elif resume.name.endswith('.docx'):
                text = extract_text_from_docx(resume)

            parsed_data = parse_resume(text)
            form = StudentProfileForm(initial=parsed_data)

        if form.is_valid():
            profile = form.save(commit=False)
            profile.user = request.user
            profile.plan = Plan.objects.last()  # Update as needed
            profile.save()
            return redirect('plans')
    else:
        form = StudentProfileForm()

    return render(request, 'profile_form.html', {'form': form})


@login_required(login_url='login')
def applied_job_create(request):
    if request.method == 'POST':
        form = AppliedJobForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            messages.success(request, 'Applied job recorded successfully.')
            return redirect('applied-job-create')  # or to a job list view
    else:
        form = AppliedJobForm()
    return render(request, 'job_form.html', {'form': form})


@login_required(login_url='login')  # protect route
def get_filtered_jobs(request):
    user = request.user
    filter_by = request.GET.get('filter', 'day')
    today = date.today()

    # default queryset
    jobs = AppliedJob.objects.none()
    today_count = week_count = month_count = 0

    if filter_by == 'day':
        jobs = AppliedJob.objects.filter(user=user, applied_date=today)
    elif filter_by == 'week':
        start_of_week = today - timedelta(days=today.weekday())
        end_of_week = start_of_week + timedelta(days=6)
        jobs = AppliedJob.objects.filter(user=user, applied_date__range=(start_of_week, end_of_week))
    elif filter_by == 'month':
        jobs = AppliedJob.objects.filter(user=user, applied_date__month=today.month, applied_date__year=today.year)

    # counts for metrics
    today_count = AppliedJob.objects.filter(user=user, applied_date=today).count()

    start_of_week = today - timedelta(days=today.weekday())
    end_of_week = start_of_week + timedelta(days=6)
    week_count = AppliedJob.objects.filter(user=user, applied_date__range=(start_of_week, end_of_week)).count()

    month_count = AppliedJob.objects.filter(user=user, applied_date__month=today.month, applied_date__year=today.year).count()

    # convert queryset to list of dicts for rendering
    records = list(jobs.values(
        'job_posting__job_title',
        'job_posting__company_name',
        'applied_date',
        'job_link'
    ))

    return render(request, 'filtered_jobs.html', {
        'records': records,
        'selected_filter': filter_by,
        'today_count': today_count,
        'week_count': week_count,
        'month_count': month_count
    })
    
    
def job_dashboard(request):
    if not request.user.is_authenticated:
        return redirect('login')  # or any other login route

    try:
        job_applier = JobApplier.objects.get(user=request.user)
        user_technologies = job_applier.technologies.all()
        matching_jobs = JobPosting.objects.filter(technologies__in=user_technologies).distinct()
    except JobApplier.DoesNotExist:
        matching_jobs = []

    context = {
        'matching_jobs': matching_jobs
    }
    return render(request, 'jobs_to_apply.html', context)